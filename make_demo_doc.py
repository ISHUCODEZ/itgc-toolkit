"""Generates ITGC_Toolkit_Demo_Speech.docx on the Desktop."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.expanduser("~"), "Desktop", "ITGC_Toolkit_Demo_Speech.docx")

YELLOW   = RGBColor(0xFF, 0xE6, 0x00)
DARK     = RGBColor(0x1A, 0x1A, 0x2E)
MUTED    = RGBColor(0x6B, 0x6B, 0x80)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT   = RGBColor(0x18, 0x8C, 0xE5)

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="IBM Plex Sans", size=11, bold=False,
             italic=False, color=DARK):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color

def shade_paragraph(para, hex_color="FFE600"):
    """Add background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)

def add_cover():
    # Yellow title bar
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    shade_paragraph(p, "1A1A2E")
    r = p.add_run("  ITGC CONTROLS TOOLKIT")
    set_font(r, size=22, bold=True, color=WHITE)

    p2 = doc.add_paragraph()
    shade_paragraph(p2, "1A1A2E")
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run("  Demo Script & Feature Overview")
    set_font(r2, size=13, color=YELLOW)

    p3 = doc.add_paragraph()
    shade_paragraph(p3, "1A1A2E")
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(14)
    r3 = p3.add_run("  Technology Risk · IT General Controls · Audit Technology")
    set_font(r3, size=9, color=RGBColor(0x9A, 0x9A, 0xA6))

def add_section_header(title, timing=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    shade_paragraph(p, "FFE600")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    label = f"  {title}"
    if timing:
        label += f"   [{timing}]"
    r = p.add_run(label)
    set_font(r, size=11, bold=True, color=DARK)

def add_body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    set_font(r, size=10.5, color=DARK)
    return p

def add_bold_body(parts, indent=False):
    """parts = list of (text, is_bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    for text, bold in parts:
        r = p.add_run(text)
        set_font(r, size=10.5, bold=bold, color=DARK)
    return p

def add_tip(text):
    p = doc.add_paragraph()
    shade_paragraph(p, "F3F3F6")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.4)
    r = p.add_run("💡  " + text)
    set_font(r, size=9.5, italic=True, color=MUTED)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("─" * 90)
    set_font(r, size=7, color=RGBColor(0xCC, 0xCC, 0xCC))

# ══════════════════════════════════════════════════════════════════════════════
add_cover()

# ── INTRO ─────────────────────────────────────────────────────────────────────
add_section_header("OPENING", "~30 seconds")
add_body(
    '"Good [morning/afternoon]. What I\'m going to show you today is an ITGC Controls Toolkit '
    'I built from scratch — a fully working audit technology platform that automates the four '
    'pillars of IT General Controls testing.'
)
add_body(
    'Most ITGC testing is still done manually — an auditor pulls a sample, pastes it into Excel, '
    'and checks boxes. This toolkit replaces that entire workflow with a deterministic, rule-based '
    'engine that tests the full population, documents every finding with a risk rationale and '
    'control reference, and monitors control health continuously over time. Let me walk you through it."'
)

# ── HOME DASHBOARD ────────────────────────────────────────────────────────────
add_section_header("THE HOME DASHBOARD", "~1 minute")
add_body(
    '"This is the home page — and unlike a typical portfolio project, it opens with live data, '
    'not static content.'
)
add_bold_body([
    ('You can see the ', False), ('Controls Reliance Opinion ', True),
    ('at the top — 83 out of 100, currently rated ', False),
    ('Reliance can be placed', True),
    ('. That\'s a real score computed from accumulated monitoring history. '
     'Below it are four live KPI tiles: open exceptions, the latest run delta '
     'showing new versus resolved findings, mean time to remediate, and total monitoring runs.', False)
])
add_bold_body([
    ('And these three cards — ', False),
    ('SoD at 80, Change at 90, Recert at 90', True),
    (' — are per-pillar health scores. So at a glance, an audit manager can see which control '
     'area is weakest before they even open a single finding.', False)
])
add_body('"Everything you\'re seeing is pulled live from the database the moment the page loads."')
add_tip("Navigate to the home page. Point to the opinion strip, the four tiles, and the three pillar score cards.")

# ── SOD ───────────────────────────────────────────────────────────────────────
add_section_header("PILLAR 1 — SoD ANALYZER", "~1.5 minutes")
add_body(
    '"Let\'s go into the first pillar — Segregation of Duties. I\'ll click Run sample export. '
    'What this does is run the full entitlements population — every user, every role — through '
    'an explicit conflict ruleset I built with eight SoD rules covering the most common fraud '
    'paths in a financial system.'
)
add_body(
    'You can see the results immediately: conflicts found, users affected, severity breakdown. '
    'Each finding cites the rule ID, the exact conflicting roles the user holds, the business '
    'risk it creates — a user who can create a vendor and approve payments can set up a fictitious '
    'vendor and pay it — and the control framework it maps to: SOX 404, COBIT DSS06.'
)
add_body(
    'That\'s how a real audit finding is documented. Not just \'user has too much access\' — '
    'but a traceable, explainable exception with risk and control context.'
)
add_body(
    'I can also click View conflict ruleset to show all eight rules — this is fully editable, '
    'so you\'d update it to match the client\'s specific application."'
)
add_tip("Click 'Run sample export' → point to a High severity finding → click 'View conflict ruleset'.")

# ── JML ───────────────────────────────────────────────────────────────────────
add_section_header("PILLAR 1 — JML ACCESS REVIEW", "~1.5 minutes")
add_body(
    '"Still in access management — this is something most firms test manually and often get wrong: '
    'Joiners, Movers and Leavers.'
)
add_body(
    'The JML control asks: when someone leaves the company, was their access revoked? It sounds '
    'simple. In practice it\'s one of the most commonly failed ITGC controls because HR records '
    'the termination and IT provisioning never gets the message.'
)
add_body(
    'I\'ll run the sample. The engine cross-references a terminated user export — name, termination '
    'date, department, reason — against the current entitlements population. It finds three former '
    'employees who still hold active system access.'
)
add_bold_body([
    ('Meera Joshi', True), (' — contract ended 60 days ago — still has DBA access and a business user role. High severity. ', False),
    ('Anita Desai', True), (' — resigned 45 days ago — still has developer, deploy-to-production, and code commit rights. Another High.', False)
])
add_body(
    'Severity is driven by how long the access has been orphaned. 30-plus days is automatically '
    'High, 7 to 30 days is Medium. The aging is evidence-based, not arbitrary."'
)
add_tip("Click 'Run sample data' on the JML page → point to the severity chips and days-since-termination.")

# ── CHANGE ────────────────────────────────────────────────────────────────────
add_section_header("PILLAR 2 — CHANGE AUDITOR", "~1 minute")
add_body(
    '"Pillar two is change management — did every change that reached production go through '
    'proper controls?'
)
add_body(
    'The engine tests three things on every change ticket: authorisation — was it approved before '
    'deployment? Testing — is there test evidence on record? And segregation — is the person who '
    'deployed the change different from the person who wrote it?'
)
add_body(
    'Each failed check is a documented exception with the change ID, the control it breaches, '
    'and the specific reason. A developer who deploys their own code to production is a classic '
    'ITGC finding — flagged automatically here. The compliance rate feeds directly into the '
    'monitoring layer and the reliance opinion score."'
)
add_tip("Click 'Run sample change log' → point to a flagged change → show the exception card.")

# ── RECERT ────────────────────────────────────────────────────────────────────
add_section_header("PILLAR 3 — RECERTIFICATION", "~1 minute")
add_body(
    '"Pillar three is access recertification — the periodic review that proves access granted '
    'months ago is still appropriate today.'
)
add_body(
    'You upload an access grant population with last-reviewed dates, set a cadence — quarterly, '
    'semi-annual, annual — and the engine classifies every grant: Current, Due Soon, Overdue, '
    'or Never Reviewed.'
)
add_body(
    'Never Reviewed is treated as High severity because there\'s no evidence the access was ever '
    'validated. Overdue is Medium. The at-risk count — those two categories combined — feeds '
    'into the reliance score."'
)
add_tip("Click 'Run sample review population' → show the status pills → point to Never Reviewed rows.")

# ── CCM ───────────────────────────────────────────────────────────────────────
add_section_header("CCM MONITOR — THE MONITORING LAYER", "~2 minutes")
add_body(
    '"This is where the toolkit goes beyond standard audit tooling — the Continuous Controls '
    'Monitoring dashboard.'
)
add_body(
    'Point-in-time testing is a snapshot — you test a sample in January and that\'s your evidence '
    'for the year. CCM is a heartbeat monitor. Every time the controls run — manually or on a '
    'scheduled cadence — the full population is tested and the result is stored as a timestamped '
    'snapshot.'
)
add_bold_body([
    ('The technically interesting part is ', False),
    ('stable finding identity', True),
    ('. Every exception gets a fingerprint — a 12-character SHA1 hash of the control name, '
     'the entity, and the rule. That means the same finding keeps its identity run to run. '
     'Which lets us do things that aren\'t possible with one-time testing.', False)
])
add_body(
    'Look at these trend charts — SoD conflicts over time, change compliance, recertification '
    'backlog. You can see the story: a spike at week six where compliance dropped to 50%, then '
    'recovery. That narrative is what you bring to a client — not just today\'s snapshot, but '
    'the trajectory.'
)
add_body(
    'The finding flow chart shows new, resolved, and persisting findings run over run. If the '
    'persisting bar is growing, management isn\'t remediating. That\'s an audit observation.'
)
add_body(
    'Below the charts is the exception register — every open finding tracked with a lifecycle '
    'status. An auditor marks it acknowledged, in-remediation, risk-accepted. There\'s an owner '
    'field and a note field. And there\'s a sign-off button — when an auditor reviews a finding, '
    'they click sign off, and their name and the date are stamped on the record. In-app audit evidence."'
)
add_tip("Show trend charts → point to week-6 spike → scroll to exception register → click 'sign off' on a finding.")

# ── OPINION + REPORTS ─────────────────────────────────────────────────────────
add_section_header("CONTROLS RELIANCE OPINION & REPORTS", "~1 minute")
add_body(
    '"At the top of the CCM page is the Controls Reliance Opinion — the conclusion an auditor '
    'reaches at the end of ITGC testing: can controls be relied upon, or does substantive testing '
    'need to expand?'
)
add_body(
    'I built this as a deterministic score: it starts at 100 and deducts for open high-severity '
    'exceptions, items open more than 90 days, weak change compliance, and a high recertification '
    'backlog. Every deduction is named — the conclusion always traces to evidence. That\'s how a '
    'real reliance memo is written.'
)
add_body(
    'From here I can generate two audit deliverables. The Excel workpaper — a seven-tab working '
    'file with a summary tab, per-control detail tabs, an exception register with sign-off columns '
    'and auto-filter, and run history. And the PDF audit report — a management-facing document '
    'with the opinion banner, control scorecard, KPI strip, and a prioritised findings table. '
    'Both are generated from live monitoring data and logged to the audit trail."'
)
add_tip("Click 'PDF Audit report' → show the download → click 'XLSX Workpaper' → show the tabs in Excel.")

# ── GOVERNANCE ────────────────────────────────────────────────────────────────
add_section_header("GOVERNANCE — AUDIT TRAIL", "~45 seconds")
add_body(
    '"Finally, the governance page. Every action in the toolkit is logged: every login, every '
    'analysis run, every export, every sign-off. You can filter by user or action type, and '
    'export the full log as a CSV.'
)
add_body(
    'The reason this matters: the toolkit practises the controls it tests. Role-based access — '
    'a viewer can read findings but cannot run controls or change exception status. An auditor '
    'can run and manage. Only an admin can modify thresholds. Every action is attributed and '
    'timestamped. The audit trail of the tool is itself audit evidence."'
)
add_tip("Filter by 'auditor' user → show sign-off and export entries → click 'CSV Export log'.")

# ── HOW WE BUILT IT ───────────────────────────────────────────────────────────
add_section_header("HOW WE BUILT IT", "~1.5 minutes")

add_bold_body([("Architecture", True)])
add_body(
    '"Flask backend, SQLite database, vanilla JavaScript frontend. No heavy framework — fast to '
    'stand up, easy to reason about. The architecture mirrors what you\'d actually deploy in a '
    'professional services context.'
)

add_bold_body([("Control Engines", True)])
add_body(
    'Four pure Python services — SoD, Change, Recert, JML — each deterministic: same input, '
    'same output, every time. The SoD engine runs eight explicit rules against a user-role matrix. '
    'The Change engine tests three controls per ticket. The Recert engine computes aging against '
    'a configurable cadence. The JML engine cross-references two populations.'
)

add_bold_body([("CCM Layer", True)])
add_body(
    'The CCM service wraps those four engines with a snapshot store, SHA1 fingerprinting, '
    'run-to-run delta detection, exception lifecycle management, threshold alerting, and '
    'APScheduler integration for automated sweeps.'
)

add_bold_body([("Reporting", True)])
add_body(
    'openpyxl for the Excel workpaper, reportlab for the PDF — both generated as in-memory '
    'byte streams and streamed to the browser via Flask\'s send_file.'
)

add_bold_body([("Security", True)])
add_body(
    'Role-based access enforced server-side on every endpoint with a require_role decorator — '
    'the frontend gates are cosmetic, the real enforcement is in the API.'
)

add_bold_body([("Demo Data", True)])
add_body(
    'A backfill script seeds 12 weeks of synthetic history through the real engines — not fake '
    'numbers, but engineered weekly datasets that produce a meaningful narrative: a compliance '
    'spike, recovery, a declining trend — so the trend charts tell a story."'
)

# ── CLOSE ─────────────────────────────────────────────────────────────────────
add_section_header("CLOSING", "~30 seconds")
add_body(
    '"To summarise — this toolkit covers every ITGC domain an auditor tests: SoD, JML, change '
    'management, recertification, continuous monitoring, and the reliance conclusion. Every '
    'finding is traceable to a named rule with risk context and a control reference. The '
    'monitoring layer converts point-in-time testing into a live heartbeat. And the output — '
    'the workpaper and the PDF report — is what you\'d actually hand to a client.'
)
add_body(
    'It\'s not a demo prototype. Every button does something real. I\'m happy to walk through '
    'any part in more detail."'
)

# ── QUICK REFERENCE TABLE ─────────────────────────────────────────────────────
add_divider()
p = doc.add_paragraph()
r = p.add_run("QUICK REFERENCE — DEMO FLOW")
set_font(r, size=9, bold=True, color=MUTED)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["Page", "What to show", "Timing"]):
    cell.text = text
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)

rows_data = [
    ("Home (/)",            "Opinion strip, KPI tiles, pillar scores",              "1 min"),
    ("SoD (/sod)",         "Run sample → conflict finding → ruleset",               "1.5 min"),
    ("JML (/jml)",         "Run sample → 3 users, severity chips",                  "1.5 min"),
    ("Change (/change)",   "Run sample → flagged change → exception card",           "1 min"),
    ("Recert (/recert)",   "Run sample → Never Reviewed rows",                       "1 min"),
    ("CCM (/ccm)",         "Trend charts → spike → exception register → sign-off",  "2 min"),
    ("CCM opinion",        "Reliance score → PDF download → XLSX download",          "1 min"),
    ("Governance",         "Filter by user → export CSV",                            "45 sec"),
    ("Build story",        "Architecture, engines, CCM, reporting, backfill",        "1.5 min"),
]
for page, what, timing in rows_data:
    row = table.add_row().cells
    row[0].text = page
    row[1].text = what
    row[2].text = timing
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.save(OUT)
print(f"Saved: {OUT}")
